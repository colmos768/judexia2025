from flask import Flask, render_template, request, redirect, url_for, Response, flash, make_response, send_from_directory, jsonify
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import text
from werkzeug.utils import secure_filename

import os
import uuid
from datetime import datetime, date
import openai
import PyPDF2
import docx
import tiktoken
import numpy as np
import io
import csv
import traceback

from models import FormatoLegal, Cliente, Causa, Documento, Honorario, PagoCuota, Contraparte, Gasto
from database import db
from dotenv import load_dotenv

ultimo_error = ""

load_dotenv()

def create_app():
    app = Flask(__name__)
    app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv("DATABASE_URL")
    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
    app.secret_key = os.getenv("FLASK_SECRET_KEY", "clave")
    app.config['UPLOAD_FOLDER'] = os.path.join("static", "documentos")

    db.init_app(app)

    os.makedirs("static/ia", exist_ok=True)
    os.makedirs("static/formatos", exist_ok=True)
    os.makedirs("static/documentos", exist_ok=True)

    openai.api_key = os.environ.get("OPENAI_API_KEY")

    with app.app_context():
        db.create_all()
   
    # =================== RUTAS GENERALES ===================

    @app.route("/")
    def index():
        return redirect(url_for("dashboard"))

    @app.route("/login")
    def login():
        return render_template("login.html")

    @app.route("/logout")
    def logout():
        return redirect(url_for("login"))

    @app.route("/dashboard")
    def dashboard():
        causas_mes = 32
        clientes_nuevos = 12
        audiencias_proximas = 8
        honorarios_pendientes = 15
        meses = ["Jan", "Feb", "Mar", "Apr", "Mai", "Jun", "Jul", "Aug"]
        grafico_causas = [3, 4, 6, 5, 9, 17, 11, 13]
        recordatorios = [
            {"texto": "Crear nuevo formato", "tag": "Hoy"},
            {"texto": "Actualizar agenda", "tag": "Hoy"},
            {"texto": "Revisar gestiones de causa", "tag": "2 días"},
            {"texto": "Programar reunión con cliente", "tag": None},
            {"texto": "Preparar informe semanal", "tag": "1 sema"}
        ]
        return render_template("dashboard.html", **locals())

    # =================== RUTAS CLIENTES ===================

    @app.route("/clientes")
    def clientes():
        lista = Cliente.query.all()
        return render_template("clientes.html", clientes=lista)

    @app.route("/registrar_cliente", methods=["POST"])
    def registrar_cliente():
        data = request.form
        nuevo = Cliente(
            nombre=data["nombre"],
            rut_num=data["rut_num"],
            rut_dv=data["rut_dv"],
            email=data.get("email"),
            telefono=data.get("telefono"),
            direccion=data.get("direccion"),
            profesion=data.get("profesion"),
            fecha_nacimiento=datetime.strptime(data["fecha_nacimiento"], "%Y-%m-%d")
        )
        db.session.add(nuevo)
        db.session.commit()
        flash("Cliente registrado correctamente.")
        return redirect(url_for("clientes"))

    # =================== RUTAS CAUSAS ===================

    @app.route("/causas", methods=["GET", "POST"])
    def causas():
        if request.method == "POST":
            data = request.form

            nueva_causa = Causa(
                tipo_causa=data["tipo_causa"],
                procedimiento=data["procedimiento"],
                judicial=data["judicial"] == "True",
                corte_apelaciones=data.get("corte_apelaciones"),
                tribunal=data.get("tribunal"),
                letra=data.get("letra"),
                rol_numero=data.get("rol_numero"),
                rol_anio=int(data.get("rol_anio") or 0),
                fecha_ingreso=datetime.strptime(data["fecha_ingreso"], "%Y-%m-%d"),
                ultima_gestion=data.get("ultima_gestion"),
                fecha_ultima_gestion=datetime.strptime(data["fecha_ultima_gestion"], "%Y-%m-%d") if data.get("fecha_ultima_gestion") else None,
                ingreso_juridico=data.get("ingreso_juridico"),
                cliente_id=int(data["cliente_id"]),
                contraparte_id=int(data["contraparte_id"]) if data.get("contraparte_id") else None
            )

            db.session.add(nueva_causa)
            db.session.commit()

            archivos = request.files.getlist("documentos")
            for archivo in archivos:
                if archivo and archivo.filename != "":
                    nombre_original = archivo.filename
                    nombre_unico = f"{uuid.uuid4().hex}_{secure_filename(nombre_original)}"
                    ruta = os.path.join("static", "documentos", nombre_unico)
                    archivo.save(ruta)

                    nuevo_doc = Documento(
                        causa_id=nueva_causa.id,
                        nombre_archivo=nombre_original,
                        ruta_archivo=ruta,
                        tipo="prueba habilitante"
                    )
                    db.session.add(nuevo_doc)

            db.session.commit()
            flash("✅ Causa registrada correctamente.")
            return redirect(url_for("causas"))

        causas = Causa.query.all()
        clientes = Cliente.query.all()
        contrapartes = Contraparte.query.all()
        return render_template("causas.html", causas=causas, clientes=clientes, contrapartes=contrapartes)

    # =================== RUTAS IA ===================

    def extraer_texto(path):
        if path.endswith(".pdf"):
            with open(path, "rb") as f:
                lector = PyPDF2.PdfReader(f)
                return " ".join([p.extract_text() or "" for p in lector.pages])
        elif path.endswith(".docx"):
            doc = docx.Document(path)
            return "\n".join([p.text for p in doc.paragraphs])
        elif path.endswith(".txt"):
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        else:
            return ""

    def dividir_en_chunks(texto, max_tokens=500):
        tokenizer = tiktoken.get_encoding("cl100k_base")
        palabras = texto.split(".")
        chunks = []
        actual = ""
        for p in palabras:
            if len(tokenizer.encode(actual + p)) < max_tokens:
                actual += p + "."
            else:
                chunks.append(actual.strip())
                actual = p + "."
        if actual:
            chunks.append(actual.strip())
        return chunks

    def obtener_embedding(texto):
        return openai.Embedding.create(
            input=texto,
            model="text-embedding-ada-002"
        )["data"][0]["embedding"]

    def similitud_coseno(a, b):
        a = np.array(a)
        b = np.array(b)
        return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))

    @app.route("/ia")
    def ia():
        archivos = [f for f in os.listdir("static/ia") if f != ".keep"]
        return render_template("ia.html", archivos=archivos)

    @app.route("/subir_ia", methods=["POST"])
    def subir_ia():
        archivo = request.files["archivo"]
        if archivo:
            ruta = os.path.join("static/ia", archivo.filename)
            archivo.save(ruta)
            flash("✅ Archivo IA subido exitosamente.")
        return redirect(url_for("ia"))

    @app.route("/eliminar_ia/<nombre>")
    def eliminar_ia(nombre):
        ruta = os.path.join("static/ia", nombre)
        if os.path.exists(ruta):
            os.remove(ruta)
            flash("🗑️ Archivo IA eliminado correctamente.")
        return redirect(url_for("ia"))

    @app.route("/preguntar_ia", methods=["POST"])
    def preguntar_ia():
        pregunta = request.form["pregunta"]
        
        archivos = sorted(
            [f for f in os.listdir("static/ia") if f.endswith((".pdf", ".docx", ".txt"))],
            key=lambda f: os.path.getmtime(os.path.join("static/ia", f)),
            reverse=True
        )
        
        if not archivos:
            flash("⚠️ No hay archivos para consultar.")
            return redirect(url_for("ia"))

        try:
            texto = extraer_texto(os.path.join("static/ia", archivos[0]))
            chunks = dividir_en_chunks(texto)
            chunks = chunks[:30]

            embeddings = [obtener_embedding(c) for c in chunks]
            pregunta_embedding = obtener_embedding(pregunta)

            top_chunk = chunks[np.argmax([similitud_coseno(e, pregunta_embedding) for e in embeddings])]

            respuesta = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Eres un asistente legal que responde en lenguaje claro."},
                    {"role": "user", "content": f"Basado en este texto: {top_chunk}\n\nResponde: {pregunta}"}
                ]
            )["choices"][0]["message"]["content"]
        except Exception as e:
            flash("❌ Error al procesar la pregunta. Intenta nuevamente más tarde.")
            print("ERROR en /preguntar_ia:", e)
            return redirect(url_for("ia"))

        archivos = [f for f in os.listdir("static/ia") if f != ".keep"]
        return render_template("ia.html", archivos=archivos, respuesta=respuesta)

    # =================== RUTAS FORMATOS ===================

    def allowed_file(filename):
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'docx', 'txt', 'jpg', 'png'}

    @app.route("/formatos", methods=["GET"])
    def formatos():
        filtro_nombre = request.args.get("nombre", "").strip()
        filtro_usuario = request.args.get("usuario", "").strip()
        filtro_causa = request.args.get("causa_id", "").strip()

        query = FormatoLegal.query

        if filtro_nombre:
            query = query.filter(FormatoLegal.nombre_original.ilike(f"%{filtro_nombre}%"))
        if filtro_usuario:
            query = query.filter(FormatoLegal.usuario.ilike(f"%{filtro_usuario}%"))
        if filtro_causa:
            query = query.filter(FormatoLegal.causa_id == int(filtro_causa))

        formatos = query.order_by(FormatoLegal.fecha_subida.desc()).all()
        causas = Causa.query.order_by(Causa.id.desc()).all()

        return render_template(
            "formatos.html",
            formatos=formatos,
            causas=causas,
            filtro_nombre=filtro_nombre,
            filtro_usuario=filtro_usuario,
            filtro_causa=filtro_causa
        )

    @app.route("/subir_formato", methods=["POST"])
    def subir_formato():
        archivo = request.files.get("archivo")
        usuario = request.form.get("usuario")
        causa_id = request.form.get("causa_id") or None
        observaciones = request.form.get("observaciones")

        if not archivo or not usuario:
            flash("Archivo y nombre de usuario son obligatorios.")
            return redirect(url_for("formatos"))

        if archivo and allowed_file(archivo.filename):
            filename = secure_filename(archivo.filename)
            path = os.path.join("static", "formatos", filename)
            archivo.save(path)

            nuevo = FormatoLegal(
                nombre_original=archivo.filename,
                filename=filename,
                usuario=usuario,
                causa_id=causa_id,
                observaciones=observaciones,
                fecha_subida=datetime.utcnow()
            )
            db.session.add(nuevo)
            db.session.commit()
            flash("✅ Formato subido correctamente.")
        else:
            flash("❌ Formato inválido. Solo se permiten archivos PDF, DOCX, TXT, JPG, PNG.")

        return redirect(url_for("formatos"))

    @app.route("/formatos/eliminar/<int:id>", methods=["POST"])
    def eliminar_formato(id):
        formato = FormatoLegal.query.get(id)
        if formato:
            ruta = os.path.join("static", "formatos", formato.filename)
            if os.path.exists(ruta):
                os.remove(ruta)
            db.session.delete(formato)
            db.session.commit()
            flash("🗑️ Formato eliminado correctamente.")
        return redirect(url_for("formatos"))

    # =================== RUTAS FACTURACIÓN ===================

    @app.route("/facturacion", methods=["GET", "POST"])
    def facturacion():
        clientes = Cliente.query.all()
        selected_cliente = request.args.get('cliente_id')
        selected_estado = request.args.get('estado')

        honorarios_query = Honorario.query
        pagos_query = PagoCuota.query

        if selected_cliente:
            honorarios_query = honorarios_query.filter_by(cliente_id=selected_cliente)
            pagos_query = pagos_query.join(Honorario).filter(Honorario.cliente_id == selected_cliente)

        if selected_estado:
            pagos_query = pagos_query.filter_by(estado=selected_estado)

        honorarios = honorarios_query.order_by(Honorario.fecha_emision.desc()).all()
        pagos = pagos_query.order_by(PagoCuota.fecha_pago.desc()).all()

        return render_template("facturacion.html", **locals())

    @app.route('/registrar_honorario', methods=['GET', 'POST'])
    def registrar_honorario():
        if request.method == 'POST':
            data = request.form
            nuevo_honorario = Honorario(
                cliente_id=data['cliente_id'],
                causa_id=data.get('causa_id') or None,
                descripcion=data['descripcion'],
                monto_total=float(data['monto_total']),
                fecha_emision=data.get('fecha_emision') or date.today(),
                en_cuotas=data.get('en_cuotas') == 'on',
                numero_cuotas=int(data.get('numero_cuotas') or 1),
                estado='pendiente'
            )
            db.session.add(nuevo_honorario)
            db.session.commit()
            return redirect(url_for('facturacion'))

        clientes = Cliente.query.all()
        causas = Causa.query.all()
        return render_template('registrar_honorario.html', clientes=clientes, causas=causas, date_today=date.today())

    @app.route('/registrar_pago/<int:honorario_id>', methods=['GET', 'POST'])
    def registrar_pago(honorario_id):
        honorario = Honorario.query.get_or_404(honorario_id)

        if request.method == 'POST':
            data = request.form
            nuevo_pago = PagoCuota(
                honorario_id=honorario.id,
                numero_cuota=int(data['numero_cuota']),
                monto_pagado=float(data['monto_pagado']),
                fecha_pago=data.get('fecha_pago') or datetime.utcnow().date(),
                vencimiento=data.get('vencimiento') or datetime.utcnow().date(),
                estado='pagado'
            )
            db.session.add(nuevo_pago)
            db.session.commit()
            return redirect(url_for('facturacion'))

        return render_template('registrar_pago.html', honorario=honorario)

    @app.route('/exportar_facturacion')
    def exportar_facturacion():
        honorarios = Honorario.query.all()
        si = io.StringIO()
        cw = csv.writer(si)
        cw.writerow(['Cliente', 'Causa', 'Descripción', 'Monto', 'Fecha', 'Cuotas', 'Estado'])

        for h in honorarios:
            cw.writerow([
                h.cliente_id,
                h.causa.rol_numero if h.causa else 'Sin causa',
                h.descripcion,
                h.monto_total,
                h.fecha_emision,
                h.numero_cuotas,
                h.estado
            ])

        output = make_response(si.getvalue())
        output.headers["Content-Disposition"] = "attachment; filename=facturacion.csv"
        output.headers["Content-type"] = "text/csv"
        return output

    # =================== RUTAS SERVICIO Y GASTOS ===================

    @app.route("/servicio")
    def servicio():
        return render_template("servicio.html")

    @app.route('/registrar_gasto', methods=['GET', 'POST'])
    def registrar_gasto():
        if request.method == 'POST':
            data = request.form
            nuevo_gasto = Gasto(
                descripcion=data['descripcion'],
                monto=float(data['monto']),
                fecha=data.get('fecha') or date.today(),
                categoria=data.get('categoria')
            )
            db.session.add(nuevo_gasto)
            db.session.commit()
            return redirect(url_for('facturacion'))

        return render_template('registrar_gasto.html', date_today=date.today())

    # =================== RUTAS DE AJUSTE Y DEPURACIÓN ===================

    @app.route("/ajustar_clientes")
    def ajustar_clientes():
        try:
            db.session.execute(text("ALTER TABLE clientes ADD COLUMN rut_num VARCHAR(8);"))
            db.session.execute(text("ALTER TABLE clientes ADD COLUMN rut_dv VARCHAR(1);"))
            db.session.execute(text("ALTER TABLE clientes ADD COLUMN profesion VARCHAR(100);"))
            db.session.commit()
            return "✅ Tabla 'clientes' actualizada correctamente."
        except Exception as e:
            return f"❌ Error: {str(e)}"

    @app.route("/initdb")
    def init_db():
        try:
            db.create_all()
            return "✅ Base de datos creada correctamente."
        except Exception as e:
            return f"❌ Error al crear la base de datos: {e}", 500

    @app.route("/debug_error")
    def debug_error():
        return f"<pre>{ultimo_error}</pre>"

    @app.errorhandler(500)
    def error_500(e):
        global ultimo_error
        ultimo_error = traceback.format_exc()
        return render_template("500.html"), 500

    # =================== RUTAS DE DEPURACIÓN Y FIX ===================

    @app.route('/ver_tablas')
    def ver_tablas():
        try:
            result = db.session.execute(text("""
                SELECT table_name 
                FROM information_schema.tables 
                WHERE table_schema = 'public'
            """))
            tablas = [row[0] for row in result]
            return '<br>'.join(tablas)
        except Exception as e:
            return f'❌ Error al listar tablas: {e}'

    @app.route('/ver_columnas_causas')
    def ver_columnas_causas():
        try:
            result = db.session.execute(text("""
                SELECT column_name
                FROM information_schema.columns
                WHERE table_name = 'causas'
            """))
            columnas = [row[0] for row in result]
            return '<br>'.join(columnas)
        except Exception as e:
            return f'❌ Error al listar columnas: {e}'

    @app.route('/fix_tipo_causa')
    def fix_tipo_causa():
        try:
            db.session.execute(text(
                "ALTER TABLE causas ADD COLUMN tipo_causa VARCHAR(100) NOT NULL DEFAULT 'Otro'"
            ))
            db.session.commit()
            return '✅ Columna tipo_causa agregada correctamente.'
        except Exception as e:
            return f'❌ Error al ejecutar ALTER TABLE: {e}'

# Al final de tu archivo app.py
app = create_app()
